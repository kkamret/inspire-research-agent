import streamlit as st
import google.generativeai as genai
from tavily import TavilyClient
from docx import Document
from io import BytesIO
from datetime import datetime

# --- 1. API 키 세팅 ---
try:
    TAVILY_API_KEY = st.secrets["TAVILY_API_KEY"]
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
except KeyError:
    st.error("API 키가 secrets.toml에 설정되지 않았습니다.")
    st.stop()

tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
genai.configure(api_key=GEMINI_API_KEY)

current_year = datetime.now().year
current_date_formatted = datetime.now().strftime('%Y-%m-%d')

# --- 2. Word 문서 생성 함수 ---
def create_word_document(title, content):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"리서치 수행일: {current_date_formatted}")
    doc.add_paragraph("\n--- 데이터 리서치 결과 ---\n")
    doc.add_paragraph(content)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- 3. 웹 UI 구성 및 디자인 ---
st.set_page_config(page_title="Marketing & IP Research Agent", layout="wide", page_icon="🔍")

inspire_style_css = """
    header, [data-testid="stDecoration"] { display: none; }
    .stApp { background-color: #F8F9FA; }
    [data-testid="stSidebar"] { background-color: #E9ECEF; border-right: 1px solid #DEE2E6; }
    
    h1, h2, h3, h4, [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] label {
        color: #4A148C !important; font-weight: 800;
    }
    
    .stButton>button {
        background-color: #4A148C; color: #FFFFFF; border: 2px solid #4A148C;
        border-radius: 6px; font-weight: bold; width: 100%; transition: 0.3s;
    }
    .stButton>button:hover {
        background-color: #E9ECEF; border-color: #4A148C; color: #4A148C;
    }
    
    .stDownloadButton>button { background-color: #FFFFFF; color: #4A148C; border: 2px solid #4A148C; }
    .stDownloadButton>button:hover { background-color: #F3E5F5; }
    
    div[data-testid="stStatusWidget"] { background-color: #F3E5F5; border: 1px solid #4A148C; color: #4A148C; }
    
    .result-container {
        background-color: #FFFFFF; padding: 30px; border-radius: 8px; 
        border: 1px solid #DEE2E6; color: #212529; box-shadow: 0px 4px 6px rgba(0,0,0,0.05);
    }
    
    .result-container table { width: 100%; }
    .result-container th, .result-container td { word-break: keep-all; }
    
    /* 탭(Tabs) 디자인 스타일링 */
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px;
    }
    .stTabs [data-baseweb="tab"] {
        background-color: #E9ECEF;
        border-radius: 6px 6px 0px 0px;
        padding: 10px 20px;
        color: #4A148C;
        font-weight: bold;
    }
    .stTabs [aria-selected="true"] {
        background-color: #4A148C;
        color: #FFFFFF !important;
    }
"""
st.markdown(f'<style>{inspire_style_css}</style>', unsafe_allow_html=True)

st.title("🔍 Marketing & IP Research Agent")
st.markdown("---")

# --- 4. 검색 조건 설정 (사이드바 - 공통 적용) ---
with st.sidebar:
    st.header("⚙️ 기본 환경 설정")
    st.markdown("*(이 설정은 모든 탭에 공통으로 적용됩니다)*")
    target_industry = st.selectbox("1. 자사/타겟 산업군", ["호텔/리조트", "카지노/복합리조트", "오프라인 유통/복합몰", "F&B/외식업"])
    specific_brand = st.text_input("2. 자사 브랜드명 (선택)", placeholder="예: 파라다이스, 더현대 등")
    
    period_options = {
        "최근 1개월": ("month", ""),
        "최근 3개월": ("year", f"{current_year}년 최근 3개월"), 
        "최근 6개월": ("year", f"{current_year}년 최근 6개월"), 
        "최근 1년": ("year", ""),
        "최근 3년": ("any", f"{current_year - 3}년 이후부터 {current_year}년까지")
    }
    selected_time = st.selectbox("3. 리서치 기간", list(period_options.keys()), index=3)
    tavily_time_range, text_time_prompt = period_options[selected_time]

# =========================================================
# --- 5. 탭(Tabs) 구성 ---
# =========================================================
tab1, tab2 = st.tabs(["📊 마케팅 트렌드 딥 리서치", "🦄 IP 콜라보 매칭 에이전트"])

# =========================================================
# [TAB 1] 기존 마케팅 트렌드 리서치 (100% 원복)
# =========================================================
with tab1:
    st.subheader("💡 리서치 키워드 조합 (모듈형)")
    st.markdown("원하는 타겟, 채널, 캠페인 방식을 자유롭게 조합하세요. AI가 최적의 검색어로 번역하여 리서치합니다.")

    col1, col2, col3 = st.columns(3)
    with col1:
        channel_pick = st.multiselect("📍 채널 (Channel)", ["오프라인 전용", "온라인/앱(App)", "온·오프라인 연계(O2O)"])
    with col2:
        target_pick = st.multiselect("🎯 타겟 고객 (Target)", ["신규 가입자", "VIP/우수고객", "휴면 고객", "2030 MZ세대", "가족 단위", "외국인 관광객"])
    with col3:
        action_pick = st.multiselect("🚀 캠페인 방식 (Action)", ["멤버십 론칭/개편", "유료 멤버십", "팝업/공간 마케팅", "할인/바우처 프로모션", "프라이빗 초청 행사", "브랜드/캐릭터 콜라보"])

    custom_query = st.text_input("✍️ 추가 자유 키워드 (선택)", placeholder="예: 두바이 초콜릿, 크리스마스 팝업, 팝업스토어 웨이팅 등 특별히 포함하고 싶은 단어")

    if st.button("Deep Research 시작 (시간 소요) 🚀", key="btn1_deep"):
        if not channel_pick and not target_pick and not action_pick and not custom_query:
            st.warning("최소 하나 이상의 키워드 블록을 선택하거나 입력해주세요.")
        else:
            with st.status(f"에이전트가 {target_industry} 트렌드를 딥 리서치 중입니다. (약 15~30초 소요)...", expanded=True) as status:
                try:
                    model = genai.GenerativeModel('gemini-2.5-flash')
                    
                    st.write("🧠 AI가 사용자의 의도를 분석하여 최적의 검색어를 설계 중입니다...")
                    
                    intent_parts = [f"산업군: {target_industry}"]
                    if specific_brand: intent_parts.append(f"타겟브랜드: {specific_brand}")
                    if channel_pick: intent_parts.append(f"채널: {', '.join(channel_pick)}")
                    if target_pick: intent_parts.append(f"타겟고객: {', '.join(target_pick)}")
                    if action_pick: intent_parts.append(f"캠페인방식: {', '.join(action_pick)}")
                    if custom_query: intent_parts.append(f"★필수 포함 추가키워드: {custom_query}")
                    
                    user_intent = "\n".join(intent_parts)
                    
                    query_gen_prompt = f"""
                    당신은 마케팅 리서치를 위한 '검색어 최적화(SEO) 전문가'입니다.
                    아래 [사용자 선택 키워드]를 참고하여, 검색 엔진에서 가장 유의미한 '대한민국 마케팅 성공 사례 기사'를 찾을 수 있는 **검색어 문자열 1줄만** 출력하십시오.
                    
                    [사용자 선택 키워드]
                    {user_intent}
                    
                    [검색어 작성 규칙 - 절대 엄수]
                    1. 과도한 따옴표 금지 (가장 중요): 모든 단어에 따옴표("")를 씌우지 마십시오. 검색 결과가 0개가 됩니다. 주요 키워드를 자연스러운 띄어쓰기로 나열하십시오.
                    2. 강력한 필수 포함: 오직 [★필수 포함 추가키워드]가 있을 경우에만, 해당 단어 하나에만 따옴표(" ")를 씌워 검색어에 포함하십시오.
                    3. 한국 기사 강제: "국내" 라는 단어를 반드시 포함할 것.
                    4. 해외 및 노이즈 차단: 문장 끝에 반드시 "-대만 -일본 -중국 -글로벌 -해외 -배달 -금융 -보험 -IT기업" 을 붙일 것.
                    5. 오직 완성된 검색어 1줄만 출력할 것.
                    """
                    
                    optimized_query = model.generate_content(query_gen_prompt).text.strip()
                    st.write(f"👉 설계된 검색어: `{optimized_query}`")

                    st.write("🔍 엄격한 필터를 적용하여 웹 데이터를 광범위하게 수집 중입니다...")
                    
                    search_params = {
                        "query": optimized_query,
                        "search_depth": "advanced",
                        "max_results": 25
                    }
                    if tavily_time_range:
                        search_params["time_range"] = tavily_time_range
                        
                    search_result = tavily_client.search(**search_params)
                    search_results_list = search_result.get('results', [])
                    
                    if not search_results_list:
                        status.update(label="검색 결과 없음", state="error", expanded=True)
                        st.error("앗! 선택하신 키워드 조합이 너무 구체적이어서 일치하는 기사를 하나도 찾지 못했습니다. 🎯 키워드 개수를 조금 줄이거나 범위를 넓혀서 다시 시도해 주세요.")
                        st.stop()
                        
                    context_text = "\n".join([f"- 제목: {res['title']}\n  내용: {res['content']}\n  출처 링크: {res['url']}\n" for res in search_results_list])
                    st.write(f"✅ {len(search_results_list)}개의 관련 문서를 수집했습니다.")

                    st.write("🧠 수집된 데이터를 깐깐하게 검증하여 최상위 사례만 표로 정리 중입니다...")
                    
                    report_prompt = f"""
                    당신은 경영진에게 보고할 '데이터 기반 전략 기획 리포트'를 작성하는 최고 수준의 깐깐한 CX/마케팅 리서처입니다.
                    아래 수집된 방대한 데이터를 바탕으로 리포트를 작성하되, **반드시 아래 규칙을 엄수**하십시오.

                    [검색된 데이터]
                    {context_text}

                    [작업 규칙 - 절대 엄수]
                    1. 해외 사례 즉각 폐기: 대한민국 영토 밖에서 벌어진 일이나 글로벌 사례는 무조건 제외하십시오.
                    2. 품질 우선주의 (Quality over Quantity): 사용자 의도에 완벽히 부합하고 내용이 구체적인 **고퀄리티 사례만 엄선하여 5개~15개 사이**로 표를 구성하십시오. 
                    3. 출처 링크 포맷 강제: 표의 '출처 링크' 칸에 긴 URL을 쓰지 마십시오. **무조건 `[기사 보기](URL)` 형태의 마크다운 하이퍼링크로만 작성**하십시오.
                    4. 이벤트 디테일 묘사: **어떤 기업이, 누구를 대상으로, 무엇을 어떻게 했는지** 상세하게 묘사하십시오.
                    5. 정량적 성과 포함: 사례 설명 시 수치적 성과가 있다면 함께 적고, 없다면 억지로 만들지 마십시오.

                    [출력 양식]
                    # 📊 마케팅 레퍼런스 딥 리서치 리포트

                    *(분석 타겟: {target_industry} / 주요 키워드: {', '.join(channel_pick + target_pick + action_pick)} {f'/ 추가 키워드: {custom_query}' if custom_query else ''})*

                    ## 1. 핵심 레퍼런스 요약 (Key Highlights)
                    - 수집된 데이터 중 가장 벤치마킹하기 좋은 구체적 성공 사례 2~3가지를 요약. (출처 링크 포함)

                    ## 2. 주요 기업별 캠페인 상세 분석 (Case Study)
                    | 기업/브랜드명 | 주요 캠페인/이벤트 상세 내용 (진행 방식 및 특징) | 확인된 성과 및 고객 반응 | 출처 링크 |
                    |---|---|---|---|
                    | ... | ... | ... | `[기사 보기](URL)` |

                    ## 3. 실무 적용을 위한 전략적 시사점 (Actionable Insights)
                    - 위 사례들을 종합하여, 우리 회사가 실제 기획 시 적용해볼 수 있는 구체적인 전략 방향 요약.
                    """
                    
                    report_content = model.generate_content(report_prompt).text
                    status.update(label="딥 리서치 및 리포트 작성 완료!", state="complete", expanded=False)
                    
                    st.subheader("📈 딥 리서치 결과 리포트")
                    st.markdown(f"<div class='result-container'>{report_content}</div>", unsafe_allow_html=True)
                    
                    docx_file = create_word_document(f"트렌드 리포트", report_content)
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.download_button(
                        label="📥 Word 문서로 다운로드",
                        data=docx_file,
                        file_name=f"Deep_Research_Report_{current_date_formatted}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_btn1"
                    )
                    
                except Exception as e:
                    status.update(label="오류 발생", state="error", expanded=True)
                    st.error(f"작업 중 에러가 발생했습니다: {e}")

# =========================================================
# [TAB 2] 신규 기능: IP 콜라보 매칭 에이전트 (물량 공세 업그레이드)
# =========================================================
with tab2:
    st.subheader("🦄 브랜드 맞춤형 IP/캐릭터 콜라보레이션 제안")
    st.markdown("원하는 콜라보 컨셉과 타겟을 선택하시면, **현재 국내에서 가장 인기 있는 수많은 IP(캐릭터, 브랜드, 아티스트 등)** 후보를 탈탈 털어 리스팅해 드립니다.")
    
    col_ip1, col_ip2 = st.columns(2)
    with col_ip1:
        ip_target = st.selectbox("🎯 메인 타겟 고객", ["전연령 (대중성)", "유아동/가족 단위", "1020 잘파세대", "2030 MZ세대", "3040 직장인/주부"], key="ip_t")
        ip_concept = st.multiselect("🎨 원하는 콜라보 컨셉 (최대 3개)", 
                                    ["귀여운/친근한", "힙한/트렌디한", "고급스러운/럭셔리", "친환경/가치소비", "유머/B급 감성", "힐링/휴식", "예술적/아티스틱"], key="ip_c")
    with col_ip2:
        ip_goal = st.selectbox("🏆 콜라보 핵심 목적", ["오프라인 공간 방문객 유도 (팝업/전시)", "신규 멤버십 가입 유도", "SNS 바이럴 및 화제성 확보", "굿즈/F&B 상품 판매 연계"], key="ip_g")

    if st.button("맞춤형 IP 대량 매칭 리서치 시작 🚀", key="btn2_ip"):
        if not ip_concept:
            st.warning("원하는 콜라보 컨셉을 최소 1개 이상 선택해주세요.")
        else:
            with st.status("최신 IP 트렌드를 분석하고 가능한 한 많은 콜라보 후보를 찾고 있습니다...", expanded=True) as status:
                try:
                    model = genai.GenerativeModel('gemini-2.5-flash')
                    
                    st.write("🧠 현재 뜨고 있는 IP 트렌드를 폭넓게 검색하기 위한 쿼리를 설계 중입니다...")
                    ip_intent = f"타겟: {ip_target}, 컨셉: {', '.join(ip_concept)}, 목적: {ip_goal}"
                    
                    ip_query_prompt = f"""
                    당신은 브랜드 IP 콜라보레이션 전문가입니다.
                    아래 사용자의 조건에 맞는 '현재 대한민국에서 유행하는 인기 캐릭터, 브랜드, 또는 아티스트 IP 콜라보 트렌드'를 찾기 위한 검색어 1줄을 작성하세요.
                    
                    [사용자 조건]
                    {ip_intent}
                    
                    [작업 규칙]
                    1. 따옴표 남발 금지. "국내" 필수 포함.
                    2. 다양한 후보를 찾기 위해 포괄적인 단어 위주로 작성 (예: 국내 2030 힙한 캐릭터 IP 콜라보레이션 인기 트렌드)
                    3. 끝에 "-대만 -일본 -중국 -글로벌 -해외" 를 반드시 붙일 것.
                    4. 검색어 1줄만 출력.
                    """
                    ip_optimized_query = model.generate_content(ip_query_prompt).text.strip()
                    st.write(f"👉 IP 검색어: `{ip_optimized_query}`")
                    
                    st.write("🔍 방대한 양의 최신 콜라보레이션 기사 및 IP 동향을 수집 중입니다...")
                    search_params_ip = {"query": ip_optimized_query, "search_depth": "advanced", "max_results": 30} # 물량 확보를 위해 30으로 상향
                    if tavily_time_range: search_params_ip["time_range"] = tavily_time_range
                        
                    ip_search_result = tavily_client.search(**search_params_ip)
                    ip_results_list = ip_search_result.get('results', [])
                    
                    if not ip_results_list:
                        status.update(label="검색 결과 없음", state="error", expanded=True)
                        st.error("관련된 IP 트렌드 기사를 찾지 못했습니다. 컨셉을 조금 더 대중적인 것으로 변경해 보세요.")
                        st.stop()
                        
                    ip_context = "\n".join([f"- 제목: {res['title']}\n  내용: {res['content']}\n  링크: {res['url']}\n" for res in ip_results_list])
                    st.write(f"✅ {len(ip_results_list)}개의 IP 트렌드 문서를 분석합니다.")
                    
                    st.write("🎯 브랜드 핏에 맞는 IP 후보들을 긁어모아 압도적인 양의 리스트를 작성 중입니다...")
                    
                    ip_report_prompt = f"""
                    당신은 {target_industry} 산업의 브랜드 기획자입니다. 우리 브랜드({specific_brand if specific_brand else '자사'})를 위한 IP 콜라보레이션 후보를 제안해야 합니다.
                    
                    [검색된 최신 IP 트렌드 데이터]
                    {ip_context}
                    
                    [사용자 요구 조건]
                    {ip_intent}
                    
                    [작업 규칙 - 절대 엄수]
                    1. 가짜 IP 창작 금지: 반드시 검색된 데이터에 등장하거나, 현재 대한민국에서 실존하는 유명 IP(캐릭터, 라이프스타일 브랜드, F&B, 일러스트레이터 등)만 추천하십시오.
                    2. 압도적인 수량(Volume): 사용자 조건과 핏이 맞는 IP 후보를 **최소 10개에서 20개 이상 최대한 많이** 발굴하여 리스팅하십시오. 절대 3~5개로 요약하거나 뭉뚱그리지 마십시오.
                    3. 아이디어 기획: 해당 IP를 우리 산업군({target_industry})에 적용했을 때의 '구체적인 공간 연출이나 프로모션 아이디어'를 각각 제안하십시오.
                    4. 출처 링크: 데이터에서 참고한 경우 `[기사 보기](URL)` 형태로 짧게 링크를 첨부하십시오.
                    
                    [출력 양식]
                    # 🦄 맞춤형 대규모 IP 콜라보레이션 매칭 리포트
                    
                    *(타겟 산업군: {target_industry} / 콜라보 목적: {ip_goal})*
                    
                    ## 1. 최신 IP 트렌드 요약
                    - 검색 데이터를 기반으로, 현재 사용자가 선택한 타겟층({ip_target})이 열광하는 콜라보 트렌드 특징 2~3줄 요약.
                    
                    ## 2. 추천 IP 후보 리스트 (Top 10+)
                    (최대한 많은 후보를 아래 표에 작성하십시오)
                    | 추천 IP (캐릭터/브랜드명) | 추천 이유 (Brand Fit) | 💡 콜라보 아이디어 (공간/이벤트 적용 방안) | 참고 링크 |
                    |---|---|---|---|
                    | ... | ... | ... | `[기사 보기](URL)` |
                    
                    ## 3. 기획자 코멘트 (Next Step)
                    - 위 수많은 후보 중 가장 강력하게 추천하는 1가지와 실무 진행 시 고려해야 할 점 1~2줄 요약.
                    """
                    
                    ip_report_content = model.generate_content(ip_report_prompt).text
                    status.update(label="대규모 IP 매칭 리포트 완성!", state="complete", expanded=False)
                    
                    st.subheader("🎯 IP 콜라보레이션 제안서")
                    st.markdown(f"<div class='result-container'>{ip_report_content}</div>", unsafe_allow_html=True)
                    
                    docx_file_ip = create_word_document("IP 콜라보 제안 리포트", ip_report_content)
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.download_button(
                        label="📥 제안서 Word 문서 다운로드", 
                        data=docx_file_ip, 
                        file_name=f"IP_Collab_Mass_Proposal_{current_date_formatted}.docx", 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                        key="dl_btn2"
                    )
                    
                except Exception as e:
                    status.update(label="오류 발생", state="error", expanded=True)
                    st.error(f"에러: {e}")
